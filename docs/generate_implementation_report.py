"""Generate GymTrack Implementation Report Word document."""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOCS = Path(__file__).resolve().parent
OUTPUT = DOCS / "GymTrack_Implementation_Report.docx"
REPO = "https://github.com/louisflim/gymtrack"

COMMITS = [
    {
        "hash": "b9449f6c12627b4f8da64460ddc13294d25aad5b",
        "short": "b9449f6",
        "date": "2026-07-01 07:45:22 +0800",
        "message": "Initial Commit",
        "description": "Project repository initialization.",
    },
    {
        "hash": "ccb74f8d23c301123722094621dc05a73332c82e",
        "short": "ccb74f8",
        "date": "2026-07-01 18:00:46 +0800",
        "message": "second commit",
        "description": "Spring Boot backend scaffold: JWT security, AuthController, User entity, login/register API, React frontend auth pages, Supabase datasource configuration.",
    },
    {
        "hash": "aa881cd7bc1ec9d2f57a29800a990d789360fa4b",
        "short": "aa881cd",
        "date": "2026-07-01 19:18:57 +0800",
        "message": "fix some bugs",
        "description": "Fixed API proxy configuration for web dev server and database connection properties.",
    },
    {
        "hash": "2b11064f1cc99f3063290e0d826c28db05249956",
        "short": "2b11064",
        "date": "2026-07-01 19:24:18 +0800",
        "message": "Create README.md",
        "description": "Added project README with SRS summary, tech stack, and setup instructions.",
    },
    {
        "hash": "6b9b568ffe249b8c187d24e38ca19b6ea2215b9e",
        "short": "6b9b568",
        "date": "2026-07-01 19:47:13 +0800",
        "message": "added dashboard",
        "description": "Renamed frontend to web/, added web dashboard page shell, included final SRS PDF in docs.",
    },
    {
        "hash": "38ec7878b0ad5b012f122f0ee28077796eec51ea",
        "short": "38ec787",
        "date": "2026-07-02 10:17:29 +0800",
        "message": "Initializes the mobile",
        "description": "Android Kotlin project setup: Retrofit API client, DataStore session, Compose auth screens, navigation graph, GymTrack theme.",
    },
    {
        "hash": "d47a6ce09f9fb42456563a940e42dedd9c9ec4d2",
        "short": "d47a6ce",
        "date": "2026-07-02 19:05:50 +0800",
        "message": "Added the mobile login and register",
        "description": "Web auth UI refactor, QR attendance backend (QrController, AttendanceController), staff creation endpoint, mobile dashboard with QR scanner and member QR card.",
    },
    {
        "hash": "284c6d0b379731278d07444e15c05e2c8db550a7",
        "short": "284c6d0",
        "date": "2026-07-04 08:05:49 +0800",
        "message": "updates",
        "description": "Major feature batch: membership/plans/payments (PayMongo), admin dashboard KPIs, member/staff management APIs, gym enrollment QR flow, full web admin panels and mobile parity.",
    },
    {
        "hash": "0e3151b367b5fd2373b1eb6f1059c430cf7a86f0",
        "short": "0e3151b",
        "date": "2026-07-05 08:20:29 +0800",
        "message": "layout updates and ui",
        "description": "Mobile UI migration from Jetpack Compose to XML layouts with Fragments, ViewBinding, RecyclerView adapters, and QrScannerDialogFragment.",
    },
]

FEATURES = [
    {
        "title": "Feature 1: User Authentication and Account Management",
        "purpose": "Secure registration and login with role-based access. Gym Owners register as Admin, Members self-register, and Admin creates Staff accounts.",
        "web": "Login/index.jsx, Register/index.jsx, RegisterForm.jsx, CreateStaffForm.jsx, api/auth.js",
        "android": "LoginFragment, RegisterFragment (XML), AuthViewModel, AuthRepository",
        "controller": "AuthController",
        "service": "AuthService",
        "repository": "UserRepository, GymRepository",
        "entity": "User, Gym, Role",
        "dto": "LoginRequest, RegisterRequest, AuthResponse, StaffAccountResponse",
        "endpoints": [
            ("POST", "/api/auth/register", "Create Admin or Member account"),
            ("POST", "/api/auth/login", "Authenticate and return JWT"),
            ("POST", "/api/auth/staff", "Admin creates Staff account"),
        ],
        "tables": "users, gyms",
        "flow": "User submits form → React/Android calls API → AuthController → AuthService validates role, hashes password (BCrypt), saves User → JwtUtil issues token → client stores session → subsequent requests include JWT via interceptor.",
        "files": [
            ("backend/.../config/SecurityConfig.java", "JWT filter, endpoint security rules"),
            ("backend/.../security/JwtAuthFilter.java", "Validates Bearer token on each request"),
            ("web/src/api/axiosInstance.js", "Attaches JWT to API calls"),
            ("mobile/.../data/remote/AuthInterceptor.kt", "Mobile JWT interceptor"),
        ],
        "testing": "Verified login/register on web and Android. Admin can create staff; staff cannot self-register. Invalid credentials return clear error messages.",
    },
    {
        "title": "Feature 2: Member Management",
        "purpose": "Admin views, searches, filters, edits members and assigns plans. Members view profile, QR code, payment and attendance history.",
        "web": "MemberTable.jsx, MemberDashboardPanel.jsx, MemberOnboardingCard.jsx, api/members.js",
        "android": "DashboardFragment (member sections), MemberAdapter, layout_member_onboarding.xml",
        "controller": "MemberController, MembershipController",
        "service": "MemberService, MembershipService",
        "repository": "UserRepository, MembershipRepository",
        "entity": "User, Membership, MembershipStatus",
        "dto": "MemberResponse, MemberUpdateRequest, AssignPlanRequest, MembershipResponse",
        "endpoints": [
            ("GET", "/api/members", "List/search/filter members (Admin)"),
            ("PUT", "/api/members/{id}", "Update member profile (Admin)"),
            ("POST", "/api/members/assign-plan", "Assign subscription plan"),
            ("GET", "/api/membership/me", "Member views own membership"),
        ],
        "tables": "users, memberships",
        "flow": "Admin opens Members tab → GET /api/members → MemberService filters by gym → table renders → Edit dialog → PUT /api/members/{id} → User row updated in Supabase.",
        "files": [
            ("backend/.../service/MemberService.java", "Member CRUD and plan assignment"),
            ("web/src/components/admin/MemberTable.jsx", "Search, status chips, edit/assign dialogs"),
        ],
        "testing": "Member list loads with search/filter. Edit member and assign plan succeed. Member dashboard shows onboarding steps and gym enrollment status.",
    },
    {
        "title": "Feature 3: Subscription Plan Management",
        "purpose": "Admin creates, edits, activates, and deactivates subscription tiers. Members see active plans for payment.",
        "web": "PlanForm.jsx, PlanList.jsx, PlanPicker.jsx, api/plans.js",
        "android": "PlanAdminAdapter, PlanAdapter, layout_plan_form.xml, item_plan.xml",
        "controller": "PlanController",
        "service": "PlanService",
        "repository": "SubscriptionPlanRepository",
        "entity": "SubscriptionPlan",
        "dto": "PlanRequest, PlanResponse",
        "endpoints": [
            ("GET", "/api/plans", "List plans (active for members, all for admin)"),
            ("POST", "/api/plans", "Create plan"),
            ("PUT", "/api/plans/{id}", "Update plan"),
        ],
        "tables": "subscription_plans",
        "flow": "Admin fills plan form → POST /api/plans → PlanService saves plan linked to gym → Member Plans tab loads active plans via GET /api/plans.",
        "files": [
            ("backend/.../service/PlanService.java", "Plan CRUD scoped to admin gym"),
            ("backend/.../util/MembershipStatusUtil.java", "Computes Active/Expiring Soon/Expired"),
        ],
        "testing": "Create/edit plan works. Inactive plans hidden from member picker. Membership status badges display correctly.",
    },
    {
        "title": "Feature 4: Online Payment (PayMongo)",
        "purpose": "Members pay for subscriptions online. Payment records stored; membership activates after payment confirmation.",
        "web": "PlanPicker.jsx, PaymentHistory.jsx, PaymentSuccess/index.jsx, api/payments.js",
        "android": "PlanAdapter, PaymentAdapter, Custom Tabs checkout in DashboardViewModel",
        "controller": "PaymentController",
        "service": "PaymentService, PayMongoService, MembershipService",
        "repository": "PaymentRepository",
        "entity": "Payment, PaymentStatus",
        "dto": "CheckoutRequest, CheckoutResponse, PaymentResponse",
        "endpoints": [
            ("POST", "/api/payments/checkout", "Create PayMongo checkout session"),
            ("POST", "/api/payments/webhook", "PayMongo payment confirmation webhook"),
            ("POST", "/api/payments/confirm-mock", "Dev mock payment confirmation"),
            ("GET", "/api/payments/me", "Member payment history"),
            ("GET", "/api/payments", "Admin views all gym payments"),
        ],
        "tables": "payments, memberships",
        "flow": "Member selects plan → POST /api/payments/checkout → PayMongoService creates session → browser opens checkout URL → on success webhook or mock confirm → PaymentService marks PAID → MembershipService.activateMembership() sets start/end dates.",
        "files": [
            ("backend/.../service/PayMongoService.java", "PayMongo API integration (GCash, Maya, card)"),
            ("backend/.../service/PaymentService.java", "Checkout, webhook handling, activation"),
        ],
        "testing": "Mock checkout flow completes and activates membership. Payment history visible on member Activity tab and admin Payments tab.",
    },
    {
        "title": "Feature 5: Attendance Tracking (QR)",
        "purpose": "QR-based gym enrollment and member check-in/check-out. Staff/Admin scan member QR; members scan gym QR to enroll.",
        "web": "QrScanner.jsx, GymQrCard.jsx, MemberQrCard.jsx, AttendanceLogTable.jsx, api/qr.js",
        "android": "QrScannerDialogFragment, layout_qr_card.xml, AttendanceLogAdapter",
        "controller": "QrController, AttendanceController",
        "service": "QrAttendanceService",
        "repository": "AttendanceLogRepository, GymRepository",
        "entity": "AttendanceLog, Gym",
        "dto": "QrCodeResponse, ScanRequest, AttendanceScanResponse, AttendanceLogResponse, MemberGymScanResponse",
        "endpoints": [
            ("GET", "/api/qr/me", "Member personal QR code"),
            ("GET", "/api/qr/gym", "Staff/Admin gym enrollment QR"),
            ("POST", "/api/attendance/scan", "Staff scans member QR (check-in/out)"),
            ("POST", "/api/attendance/scan-gym", "Member scans gym QR to enroll"),
            ("GET", "/api/attendance/me", "Member attendance history"),
            ("GET", "/api/attendance/gym", "Admin attendance logs with filters"),
        ],
        "tables": "attendance_logs, users, gyms",
        "flow": "Staff opens scanner → camera reads QR payload → POST /api/attendance/scan → QrAttendanceService validates membership → creates/updates AttendanceLog (check_in_time, check_out_time) → returns member name, plan, status, action.",
        "files": [
            ("backend/.../service/QrAttendanceService.java", "QR generation, scan logic, enrollment"),
            ("web/src/components/qr/QrScanner.jsx", "Webcam QR scanner with camera cleanup"),
        ],
        "testing": "Gym QR enrollment works. Member QR scan toggles check-in/out. Expired memberships denied. Attendance logs filterable by date and name.",
    },
    {
        "title": "Feature 6: Admin Dashboard",
        "purpose": "Admin overview of gym KPIs, staff management, and consolidated admin tabs for plans, members, staff, attendance, and payments.",
        "web": "AdminDashboardPanel.jsx, KpiSummaryGrid.jsx, Dashboard/index.jsx, api/dashboard.js",
        "android": "DashboardFragment admin sections, KpiAdapter, layout_create_staff.xml",
        "controller": "DashboardController, StaffController",
        "service": "DashboardService, StaffService",
        "repository": "UserRepository, PaymentRepository, AttendanceLogRepository",
        "entity": "User (staff/members), Payment, AttendanceLog",
        "dto": "DashboardStatsResponse, StaffResponse, StaffUpdateRequest",
        "endpoints": [
            ("GET", "/api/dashboard/stats", "KPI metrics for admin home"),
            ("GET", "/api/staff", "List staff accounts"),
            ("PUT", "/api/staff/{id}", "Edit/deactivate staff"),
        ],
        "tables": "users, memberships, payments, attendance_logs",
        "flow": "Admin opens Home tab → GET /api/dashboard/stats → DashboardService aggregates members, active subs, expired, today check-ins, total payments → KPI grid renders.",
        "files": [
            ("backend/.../service/DashboardService.java", "Aggregates gym statistics"),
            ("web/src/components/admin/AdminDashboardPanel.jsx", "Tabbed admin UI"),
        ],
        "testing": "KPI cards populate with live data. Staff create/edit works from Overview tab. All admin tabs load without errors on web and mobile.",
    },
]


def set_cell_shading(cell, fill: str):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def add_title(doc, text, level=0):
    if level == 0:
        p = doc.add_heading(text, level=0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_heading(text, level=level)


def add_table(doc, headers, rows, header_fill="1F2937"):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_fill)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()


def build_doc():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_title(doc, "GymTrack — Implementation Documentation")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        f"Louis Francis Lim | IT342-G01 | Generated {datetime.now().strftime('%B %d, %Y')}"
    )
    run.italic = True

    doc.add_page_break()

    # Section 1
    add_title(doc, "1. GitHub Commit History", level=1)
    doc.add_paragraph(f"Repository: {REPO}")
    doc.add_paragraph(
        "Commits page: https://github.com/louisflim/gymtrack/commits/main"
    )
    doc.add_paragraph(
        "Note: Insert a screenshot of the GitHub commit history page in the space below "
        "(GitHub → gymtrack repository → Commits tab)."
    )
    doc.add_paragraph("[ SCREENSHOT PLACEHOLDER — GitHub Commit History ]").italic = True
    doc.add_paragraph()

    add_table(
        doc,
        ["#", "Commit Hash", "Date", "Commit Message", "Description"],
        [
            [
                str(i + 1),
                c["short"] + "\n" + c["hash"],
                c["date"],
                c["message"],
                c["description"],
            ]
            for i, c in enumerate(COMMITS)
        ],
    )

    doc.add_heading("Commit-to-Feature Mapping", level=2)
    mapping = [
        ("User Authentication (backend + web)", "ccb74f8"),
        ("Web dashboard shell", "6b9b568"),
        ("Android project initialization", "38ec787"),
        ("QR attendance + web/mobile auth UI", "d47a6ce"),
        ("Plans, payments, membership, admin dashboard", "284c6d0"),
        ("Mobile XML UI migration", "0e3151b"),
    ]
    add_table(
        doc,
        ["Feature Area", "Primary Commit"],
        [[m[0], m[1]] for m in mapping],
    )

    doc.add_paragraph(
        "Submission note: Commits 284c6d0 and 0e3151b bundle multiple features. "
        "For future work, split each SRS feature into its own commit with a descriptive message "
        "(e.g., feat(backend): add PayMongo checkout endpoint)."
    )

    doc.add_page_break()

    # Section 2
    add_title(doc, "2. Implementation Explanation Report", level=1)

    for feat in FEATURES:
        doc.add_heading(feat["title"], level=2)
        doc.add_paragraph(f"Purpose: {feat['purpose']}")
        doc.add_paragraph(f"ReactJS components: {feat['web']}")
        doc.add_paragraph(f"Android screens/fragments: {feat['android']}")
        doc.add_paragraph(
            f"Spring Boot — Controller: {feat['controller']} | "
            f"Service: {feat['service']} | Repository: {feat['repository']} | "
            f"Entity: {feat['entity']} | DTO: {feat['dto']}"
        )
        doc.add_paragraph("API endpoints:")
        add_table(
            doc,
            ["HTTP Method", "Endpoint", "Description"],
            [[e[0], e[1], e[2]] for e in feat["endpoints"]],
        )
        doc.add_paragraph(f"Database tables: {feat['tables']}")
        doc.add_paragraph(f"Data flow: {feat['flow']}")
        doc.add_paragraph("Key source files:")
        for path, role in feat["files"]:
            doc.add_paragraph(f"• {path} — {role}", style="List Bullet")
        doc.add_paragraph(f"Testing result: {feat['testing']}")
        doc.add_paragraph(
            "[ SCREENSHOT PLACEHOLDER — insert UI test screenshot for this feature ]"
        ).italic = True
        doc.add_paragraph()

    doc.add_page_break()

    # Section 3
    add_title(doc, "3. System Architecture Diagram", level=1)

    doc.add_paragraph(
        "GymTrack uses a three-tier architecture: ReactJS web app and Android Kotlin app "
        "both consume the same Spring Boot REST API, which persists data in Supabase (PostgreSQL) "
        "and integrates PayMongo for online payments."
    )

    arch_rows = [
        ("ReactJS Web App", "Browser UI", "axios → REST API", "Login, Dashboard, QR scanner, Admin panels"),
        ("Android Kotlin App", "Mobile UI", "Retrofit → REST API", "Fragments, ViewBinding, ZXing QR scanner"),
        ("Spring Boot Backend", "REST API layer", "Controllers → Services → Repositories", "JWT auth, business logic"),
        ("Supabase PostgreSQL", "Database", "JDBC / Hibernate JPA", "users, gyms, plans, memberships, payments, attendance_logs"),
        ("PayMongo", "External payment gateway", "HTTPS webhook + checkout API", "GCash, Maya, credit/debit card"),
    ]
    add_table(
        doc,
        ["Component", "Layer", "Connection", "Responsibility"],
        arch_rows,
    )

    doc.add_heading("Architecture Diagram", level=2)
    diagram_png = DOCS / "architecture_diagram.png"
    if diagram_png.exists():
        doc.add_picture(str(diagram_png), width=Inches(6.0))
        doc.add_paragraph()

    diagram = """
┌─────────────────────┐       ┌─────────────────────┐
│   ReactJS Web App   │       │ Android Kotlin App  │
│  (Chrome/Firefox)   │       │  (Retrofit + ZXing) │
└──────────┬──────────┘       └──────────┬──────────┘
           │         HTTPS / JSON         │
           └──────────────┬───────────────┘
                          ▼
              ┌───────────────────────┐
              │  Spring Boot Backend  │
              │  REST API + JWT Auth  │
              │  Port 8080            │
              └───────────┬───────────┘
                          │
           ┌──────────────┼──────────────┐
           ▼              ▼              ▼
   ┌──────────────┐ ┌──────────┐ ┌─────────────┐
   │  Supabase    │ │ PayMongo │ │  ZXing QR   │
   │  PostgreSQL  │ │ Checkout │ │  (embedded) │
   └──────────────┘ └──────────┘ └─────────────┘
"""
    p = doc.add_paragraph()
    run = p.add_run(diagram)
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.add_heading("Request Flow (Example: Member Payment)", level=2)
    flow_steps = [
        "1. Member taps Subscribe on web or Android Plans tab.",
        "2. Client sends POST /api/payments/checkout with planId and JWT.",
        "3. PaymentController → PaymentService creates PENDING payment row.",
        "4. PayMongoService creates checkout session; returns checkout URL.",
        "5. Client opens PayMongo hosted page (GCash/Maya/card).",
        "6. PayMongo sends POST /api/payments/webhook on success.",
        "7. PaymentService marks payment PAID; MembershipService activates membership.",
        "8. Member dashboard refreshes; status shows Active with new end date.",
    ]
    for step in flow_steps:
        doc.add_paragraph(step, style="List Number")

    doc.add_heading("Request Flow (Example: QR Check-in)", level=2)
    qr_steps = [
        "1. Staff opens QR scanner on web or mobile Scan tab.",
        "2. Camera reads member QR payload (gymtrack:member:{userId}).",
        "3. Client sends POST /api/attendance/scan with qrData and JWT.",
        "4. QrAttendanceService validates membership status (Active/Expiring).",
        "5. If no open visit: creates attendance_logs row with check_in_time.",
        "6. If open visit exists: sets check_out_time on same row.",
        "7. Response shows member name, plan, status, and CHECK_IN or CHECK_OUT.",
    ]
    for step in qr_steps:
        doc.add_paragraph(step, style="List Number")

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    build_doc()
